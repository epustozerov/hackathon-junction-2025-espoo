import os
import tempfile
import re
import yaml
from openai import OpenAI
import sys
import markdown
from html.parser import HTMLParser

try:
    from config.config import OPENAI_API_KEY
except ImportError:
    OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')

if not OPENAI_API_KEY:
    print("Warning: OPENAI_API_KEY not found. Please set it in config/config.py or as an environment variable.")
    sys.exit(1)

client = OpenAI(api_key=OPENAI_API_KEY)


def load_yaml_answers(yaml_path):
    with open(yaml_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    answers = {}
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        question_match = re.match(r'"([^"]+)":', line)
        if question_match:
            question_label = question_match.group(1)
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if next_line.startswith('answer:'):
                    answer = next_line.replace('answer:', '').strip()
                    if answer == 'None' or answer == '':
                        break
                    if answer.startswith('"') and answer.endswith('"'):
                        answer = answer[1:-1]
                        answer = answer.replace('\\"', '"').replace('\\n', '\n')
                    if answer and answer.strip():
                        answers[question_label] = answer.strip()
                    break
                elif next_line and not next_line.startswith('#') and ':' in next_line:
                    break
                i += 1
        i += 1
    
    return answers


def build_filling_prompt(template_markdown, answers):
    import yaml
    yaml_text = yaml.safe_dump(answers, sort_keys=False, allow_unicode=True)
    parts = []
    parts.append(
        "You are helping to fill in a business plan document.\n"
        "You will receive:\n"
        "1) A markdown version of a business plan template.\n"
        "2) A YAML structure containing questions and their answers.\n\n"
        "CRITICAL INSTRUCTIONS:\n"
        "- ONLY include sections and questions that have answers in the YAML.\n"
        "- COMPLETELY REMOVE any sections, questions, or placeholders that do not have answers.\n"
        "- Do NOT include '...' placeholders or empty sections.\n"
        "- If a section has no answered questions, remove the entire section.\n"
        "- If a question has no answer, remove that question entirely.\n"
        "- Preserve the structure and formatting of sections that DO have answers.\n"
        "- Insert the YAML answers into the appropriate sections of the template.\n"
        "- Do not add explanations of what you are doing; return only the final filled markdown.\n\n"
    )
    parts.append("Here is the template in markdown:\n\n```markdown\n")
    parts.append(template_markdown)
    parts.append("\n```\n\n")
    parts.append("Here are the answers in YAML (ONLY use questions that appear here):\n\n```yaml\n")
    parts.append(yaml_text)
    parts.append("\n```\n\n")
    parts.append(
        "Please return the business plan as markdown, including ONLY sections and questions that have answers.\n"
        "\n"
        "Important formatting instructions:\n"
        "- Add blank lines before and after all headings, lists, and tables.\n"
        "- Remove all <br> tags inside tables; replace them with '-' or paragraph formatting as appropriate.\n"
        "- Ensure nested lists use consistent indentation.\n"
        "- Do not insert manual line breaks within paragraphs; let Markdown handle text wrapping naturally.\n"
        "- Remember: If a question or section has no answer in the YAML, do NOT include it in the output.\n"
    )
    return "".join(parts)


def fill_business_plan_markdown_from_answers(template_path, answers):
    with open(template_path, "r", encoding="utf-8") as f:
        template_markdown = f.read()
    
    if not answers:
        raise ValueError("No answers provided. Please answer some questions first.")
    
    print(f"Using {len(answers)} answers: {list(answers.keys())}")
    
    prompt = build_filling_prompt(template_markdown, answers)
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {'role': 'system', 'content': 'You are a helpful assistant that fills business plan templates with provided answers.'},
                {'role': 'user', 'content': prompt}
            ],
            temperature=0.3
        )
        filled_markdown = response.choices[0].message.content.strip()
        return filled_markdown
    except Exception as e:
        print(f"Error filling business plan: {str(e)}")
        raise


class HTMLToDocxParser(HTMLParser):
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.list_level = 0
        self.in_list = False
        
    def handle_starttag(self, tag, attrs):
        if tag == 'h1':
            self.current_paragraph = self.doc.add_heading(level=1)
            self.current_run = None
        elif tag == 'h2':
            self.current_paragraph = self.doc.add_heading(level=2)
            self.current_run = None
        elif tag == 'h3':
            self.current_paragraph = self.doc.add_heading(level=3)
            self.current_run = None
        elif tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
            self.current_run = None
        elif tag == 'strong' or tag == 'b':
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
            self.current_run = self.current_paragraph.add_run()
            self.current_run.bold = True
        elif tag == 'em' or tag == 'i':
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
            self.current_run = self.current_paragraph.add_run()
            self.current_run.italic = True
        elif tag == 'ul':
            self.in_list = True
            self.list_level += 1
        elif tag == 'ol':
            self.in_list = True
            self.list_level += 1
        elif tag == 'li':
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph(style='List Bullet' if self.list_level == 1 else f'List Bullet {self.list_level}')
            self.current_run = None
        elif tag == 'br':
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.add_run().add_break()
        elif tag == 'hr':
            self.current_paragraph = self.doc.add_paragraph('_' * 50)
            self.current_run = None
            
    def handle_endtag(self, tag):
        if tag in ['h1', 'h2', 'h3', 'p']:
            self.current_paragraph = None
            self.current_run = None
        elif tag in ['ul', 'ol']:
            self.list_level = max(0, self.list_level - 1)
            if self.list_level == 0:
                self.in_list = False
        elif tag in ['strong', 'b', 'em', 'i']:
            self.current_run = None
            
    def handle_data(self, data):
        if data.strip():
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
            if self.current_run is None:
                self.current_run = self.current_paragraph.add_run()
            self.current_run.text += data


def create_docx_from_form_data(form_data, business_plan_sections, output_docx_path=None):
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    template_path = os.path.join(base_dir, 'business_plan', 'business_plan_template.md')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    if not form_data:
        raise ValueError("Form data is empty. Please start a conversation and answer some questions first.")
    
    print(f"DEBUG: form_data type: {type(form_data)}, keys: {list(form_data.keys()) if form_data else 'None'}")
    print(f"DEBUG: form_data items: {[(k, str(v)[:50] if v else None) for k, v in (form_data.items() if form_data else [])]}")
    print(f"DEBUG: business_plan_sections count: {len(business_plan_sections) if business_plan_sections else 0}")
    
    answers = {}
    total_questions_checked = 0
    question_ids_looked_for = []
    
    initial_form_mapping = {
        'company_name': 'Company Name',
        'sphere': 'Business Sphere / Industry',
        'location': 'Location',
        'education': 'Education',
        'experience': 'Experience / Background'
    }
    
    for key, label in initial_form_mapping.items():
        value = form_data.get(key)
        if value and isinstance(value, str) and value.strip() and value != '':
            answers[label] = value.strip()
            print(f"DEBUG: Added initial form field '{label}': {value[:50]}")
    
    for section in business_plan_sections:
        for question in section.get('core_questions', []):
            total_questions_checked += 1
            question_id = question.get('id')
            question_label = question.get('label')
            if not question_id:
                continue
            question_ids_looked_for.append(question_id)
            answer = form_data.get(question_id)
            print(f"DEBUG: Looking for question_id='{question_id}', label='{question_label}', found in form_data: {question_id in form_data}, value: {str(answer)[:50] if answer else 'None'}")
            if answer and isinstance(answer, str) and answer.strip() and answer != '':
                answers[question_label] = answer.strip()
                print(f"DEBUG: ✓ Found answer for '{question_label}': {answer[:50]}")
        for question in section.get('optional_questions', []):
            total_questions_checked += 1
            question_id = question.get('id')
            question_label = question.get('label')
            if not question_id:
                continue
            question_ids_looked_for.append(question_id)
            answer = form_data.get(question_id)
            print(f"DEBUG: Looking for optional question_id='{question_id}', label='{question_label}', found in form_data: {question_id in form_data}, value: {str(answer)[:50] if answer else 'None'}")
            if answer and isinstance(answer, str) and answer.strip() and answer != '':
                answers[question_label] = answer.strip()
                print(f"DEBUG: ✓ Found optional answer for '{question_label}': {answer[:50]}")
    
    print(f"DEBUG: Question IDs we looked for: {question_ids_looked_for[:10]}... (showing first 10)")
    print(f"DEBUG: Form data keys: {list(form_data.keys())}")
    print(f"DEBUG: Matching keys: {set(question_ids_looked_for) & set(form_data.keys())}")
    
    print(f"DEBUG: Checked {total_questions_checked} questions, found {len(answers)} answers")
    print(f"DEBUG: Final answers dict keys: {list(answers.keys())}")
    
    if not answers:
        available_keys = list(form_data.keys()) if form_data else []
        initial_form_keys = ['company_name', 'language', 'sphere', 'education', 'experience', 'location']
        has_initial_form = any(key in form_data for key in initial_form_keys)
        has_business_plan_answers = any(key not in initial_form_keys + ['email', 'report_sent'] for key in form_data.keys())
        
        if not has_initial_form:
            raise ValueError("Please complete the initial form first (company name, language, business sphere, education, experience, and location).")
        elif not has_business_plan_answers:
            raise ValueError("Please answer some business plan questions first. Complete the initial form, then the system will ask you business plan questions.")
        else:
            raise ValueError(f"No valid business plan answers found. Form data has keys: {available_keys}. Please make sure you've answered some business plan questions.")
    
    print(f"Loaded {len(answers)} answers from form_data: {list(answers.keys())}")
    
    filled_markdown = fill_business_plan_markdown_from_answers(template_path, answers)
    
    if output_docx_path is None:
        temp_dir = tempfile.gettempdir()
        output_docx_path = os.path.join(temp_dir, f'business_plan_{os.getpid()}.docx')
    
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        md = markdown.Markdown(extensions=['extra', 'tables', 'nl2br'])
        html_content = md.convert(filled_markdown)
        
        parser = HTMLToDocxParser(doc)
        parser.feed(html_content)
        
        doc.save(output_docx_path)
        
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX generation. "
            "Install it with: pip install python-docx markdown"
        )
    except Exception as e:
        print(f"Error creating DOCX: {str(e)}")
        raise
    
    return output_docx_path

